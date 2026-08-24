"""Generate two hand-off .docx documents.

1. Manager quick-start (cover-safe: connection only, zero hidden functionality).
2. DevOps deploy notes (concise, assumes competence; .env specifics + sizing).

Run:
    .venv/Scripts/python.exe -m scripts.gen_handoff_docs
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
MANAGER_OUT = ROOT / "Partner Assistant — инструкция для менеджера.docx"
DEVOPS_OUT = ROOT / "Partner Assistant — деплой (DevOps).docx"

ACCENT = RGBColor(0x3B, 0x4C, 0xC0)   # indigo
WARN = RGBColor(0xB9, 0x1C, 0x1C)     # red
MUTED = RGBColor(0x70, 0x70, 0x70)


# --------------------------------------------------------------------------- #
# Helpers
# --------------------------------------------------------------------------- #
def title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = ACCENT


def subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = MUTED


def h1(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = ACCENT


def h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def para(doc: Document, text: str = "") -> None:
    doc.add_paragraph(text)


def warn(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"❗ {text}")
    run.font.bold = True
    run.font.color.rgb = WARN


def bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def numbered(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Number")


def mono(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, hcell in enumerate(headers):
        run = t.rows[0].cells[i].paragraphs[0].add_run(hcell)
        run.font.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


# --------------------------------------------------------------------------- #
# Document 1 — Manager quick start
# --------------------------------------------------------------------------- #
def build_manager() -> None:
    doc = Document()
    title(doc, "Partner Assistant — подключение")
    subtitle(doc, "Краткая инструкция для менеджера")
    para(doc)
    para(
        doc,
        "Partner Assistant — бот-помощник для работы с партнёрскими чатами. "
        "Ниже три шага, чтобы его подключить.",
    )

    h1(doc, "1. Регистрация (один раз)")
    numbered(
        doc,
        [
            "Открой личный чат с ботом и отправь команду: /register",
            "Пришли свой Slack member ID (Slack → клик на аватар → View Profile "
            "→ More (…) → Copy member ID; выглядит как U01234ABCDE).",
            "Введи код, который придёт тебе в Slack. Готово — ты подключён и "
            "будешь получать уведомления.",
        ],
    )
    para(doc, "По желанию укажи рабочие часы, чтобы уведомления приходили в твою смену:")
    mono(doc, "/set_hours 09:00-18:00 Europe/Kiev")

    h1(doc, "2. Добавить бота в партнёрскую группу")
    numbered(
        doc,
        [
            "Назови группу в формате:  affiliate_id | Партнёр | Beton.Win",
            "Добавь бота в группу обычным участником.",
            "Всё — группа подключается автоматически, ничего больше делать не нужно.",
        ],
    )
    para(doc, "Пример названия группы:")
    mono(doc, "78516 | Acme Media | Beton.Win")

    h1(doc, "3. Secretary mode — личные диалоги с партнёром (по желанию)")
    numbered(
        doc,
        [
            "Telegram → Настройки → Telegram for Business → Чат-боты.",
            "Выбери этого бота и разреши ему доступ к сообщениям (read messages).",
            "После включения доступ активируется. Если не заработало сразу — "
            "напиши администратору, он подтвердит подключение.",
        ],
    )

    h1(doc, "Что видит бот")
    bullets(
        doc,
        [
            "Сообщения (текст, голос, файлы, пересланное) в подключённых группах "
            "и — в secretary mode — в твоих личных диалогах с партнёрами.",
            "Бот НЕ видит твои остальные личные чаты.",
            "Бот ничего не пишет в партнёрских чатах — он только помощник на твоей стороне.",
        ],
    )

    doc.save(str(MANAGER_OUT))
    print(f"wrote: {MANAGER_OUT.name}")


# --------------------------------------------------------------------------- #
# Document 2 — DevOps deploy notes
# --------------------------------------------------------------------------- #
def build_devops() -> None:
    doc = Document()
    title(doc, "Partner Assistant — заметки по деплою")
    subtitle(doc, "Для DevOps. Только особенности системы — без базовых вещей.")
    para(doc)

    h1(doc, "1. Доступ к репозиторию (приватный)")
    bullets(
        doc,
        [
            "Репозиторий: warden-afk/tg_ai_bot_bow (GitHub, private).",
            "Поделиться: GitHub → репо → Settings → Collaborators and teams → "
            "Invite — пригласить аккаунт DevOps (доступ Read достаточно).",
            "Для CI/сервера без личного аккаунта — Deploy key (read-only SSH) в "
            "Settings → Deploy keys.",
            "Текущий прод: Railway, подключён к этому репо — авто-деплой на push в main.",
        ],
    )

    h1(doc, "2. Стек и архитектура (как есть сейчас)")
    bullets(
        doc,
        [
            "Python 3.11, FastAPI + uvicorn — один процесс, порт 8080.",
            "Telegram — webhook (aiogram). На том же процессе: Slack-callbacks, "
            "веб-страницы отчётов, /health.",
            "БД — Supabase Postgres через asyncpg (пул 2–10). Ни Redis, ни брокера, "
            "ни Celery — очередь задач в самой БД.",
            "Фоновые воркеры — in-process asyncio (анализ, файлы, whisper-drain, "
            "summary-scheduler, ops-alerts, reapers). Стартуют в FastAPI lifespan, "
            "работают постоянно между запросами.",
            "LLM — OpenRouter (Haiku — анализ, Sonnet — саммари). Whisper (OpenAI) "
            "выключен флагом.",
            "Docker: python:3.11-slim + ffmpeg + curl; ставит из pyproject; non-root "
            "(uid 1000); HEALTHCHECK на /health.",
        ],
    )

    h1(doc, "3. Критично при деплое")
    warn(
        doc,
        "Никакого scale-to-zero / auto-stop на idle. Воркеры крутятся постоянно "
        "между HTTP-запросами; если платформа усыпляет инстанс — плановые рассылки "
        "(отчёты, ops-alerts) молча перестанут работать. На Fly: "
        "auto_stop_machines=false, min_machines_running=1.",
    )
    warn(
        doc,
        "Один инстанс. Воркеры in-process; горизонтальное масштабирование "
        "задублирует плановые рассылки (часть защищена dedup в БД, но не всё). "
        "Держать ровно одну реплику.",
    )
    bullets(
        doc,
        [
            "Webhook регистрируется при каждом старте, URL = SERVER_BASE_URL + "
            "/webhook. Нужен публичный HTTPS с валидным сертификатом (Railway/Fly "
            "дают из коробки; на голом сервере — reverse-proxy + TLS).",
            "Миграции НЕ применяются автоматически. supabase/migrations/0001–0016 "
            "прогнать вручную (Supabase SQL Editor или скриптом). Таблица "
            "schema_migrations не ведётся — ориентируйся по факту.",
            "/health отдаёт 503, если БД недоступна — годится для healthcheck/LB.",
        ],
    )

    h1(doc, "4. Особенности .env")
    bullets(
        doc,
        [
            "SERVER_BASE_URL — единственный URL, который надо задать. Webhook "
            "выводится из него автоматически; TELEGRAM_WEBHOOK_URL задавать НЕ нужно.",
            "SUPABASE_DB_URL — строка подключения asyncpg. Если в пароле есть "
            "спецсимволы (например +), задать SUPABASE_DB_PASSWORD отдельно — он "
            "перебивает пароль из URL.",
            "SUMMARY_ACCESS_TOKEN — bearer для ручного POST /summary/generate. К "
            "доступу к самим отчётам отношения не имеет (у отчётов свои share-токены "
            "+ пароль).",
            "Slack: SLACK_BOT_TOKEN, SLACK_SIGNING_SECRET, SLACK_CHANNEL_ALERTS, "
            "SLACK_CHANNEL_REPORTS.",
            "DAILY_LLM_BUDGET_USD (по умолчанию 30) — дневной circuit breaker на LLM.",
            "Killswitches: WHISPER_ENABLED (false), FILE_ANALYSIS_ENABLED (true).",
            "Ops-alerts: OPS_ALERTS_ENABLED + OPS_FEED_URL. OPS_FEED_URL "
            "чувствительный — только в окружении деплоя, в коде/доках его нет.",
            "Секреты — только через env. Никаких .env в репозитории.",
        ],
    )
    para(doc, "Обязательные (без дефолтов — без них процесс не стартует):")
    mono(
        doc,
        "TELEGRAM_BOT_TOKEN  TELEGRAM_WEBHOOK_SECRET  SUPABASE_URL  "
        "SUPABASE_SERVICE_KEY\nSUPABASE_DB_URL  OPENROUTER_API_KEY  OPENAI_API_KEY  "
        "SLACK_BOT_TOKEN\nSLACK_SIGNING_SECRET  SLACK_CHANNEL_ALERTS  "
        "SLACK_CHANNEL_REPORTS",
    )
    para(
        doc,
        "Плюс на проде обязательно переопределить SERVER_BASE_URL и "
        "SUMMARY_ACCESS_TOKEN (у них есть дефолты-заглушки).",
    )

    h1(doc, "5. Требования к серверу (ориентир, без запаса «на всякий»)")
    table(
        doc,
        ["Ресурс", "Значение", "Почему"],
        [
            ["CPU", "1 vCPU (shared ОК), 2 — комфортно",
             "один процесс, нагрузка I/O-bound (HTTP к LLM/Telegram/БД)"],
            ["RAM", "512 MB рабочих, 1 GB рекомендовано",
             "Python + пул соединений + HTTP-буферы + извлечение текста из файлов "
             "(до ~40k симв) + загрузка файлов до 20 MB + ffmpeg; Whisper выключен"],
            ["Диск", "~5 GB",
             "образ + зависимости; постоянного состояния на диске нет — всё в Supabase"],
            ["Сеть", "исходящий + входящий HTTPS",
             "наружу: Telegram, Supabase, OpenRouter, OpenAI, Slack, ops-фид; "
             "внутрь: webhook с валидным TLS"],
            ["Runtime", "Python 3.11+, ffmpeg",
             "ffmpeg нужен для извлечения аудио из video_note"],
        ],
    )

    doc.save(str(DEVOPS_OUT))
    print(f"wrote: {DEVOPS_OUT.name}")


if __name__ == "__main__":
    build_manager()
    build_devops()
